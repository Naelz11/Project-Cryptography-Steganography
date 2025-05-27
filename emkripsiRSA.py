# -*- coding: utf-8 -*-

from math import e
from PyQt5 import QtCore, QtGui, QtWidgets
from Crypto.PublicKey import RSA
from Crypto.Cipher import PKCS1_OAEP
import base64
import docx
import time

from startwindow import Ui_startwindow

class Ui_Enkripsi(object):
    def setupUi(self,Enkripsi,public_key_n=None, public_key_e=None):
        Enkripsi.setObjectName("Enkripsi")
        Enkripsi.resize(599, 698)
        Enkripsi.setStyleSheet("")
        
        self.label_2 = QtWidgets.QLabel(Enkripsi)
        self.label_2.setGeometry(QtCore.QRect(100, 230, 71, 16))
        self.label_2.setStyleSheet("background-color:rgb(2, 36, 64); color:rgb(255, 255, 255); font-size:12pt;")
        self.label_2.setObjectName("label_2")
        
        self.inputdata = QtWidgets.QLineEdit(Enkripsi)
        self.inputdata.setGeometry(QtCore.QRect(190, 220, 201, 31))
        self.inputdata.setStyleSheet("color:rgb(255, 255, 255); background-color:rgb(2, 36, 64); font-size:11pt;")
        self.inputdata.setObjectName("inputdata")
        
        self.browsebutton = QtWidgets.QPushButton(Enkripsi)
        self.browsebutton.setGeometry(QtCore.QRect(420, 220, 75, 31))
        self.browsebutton.setStyleSheet("background-color:rgb(255, 255, 255); border-radius: 5px;")
        self.browsebutton.setObjectName("browsebutton")
        
        self.label_3 = QtWidgets.QLabel(Enkripsi)
        self.label_3.setGeometry(QtCore.QRect(100, 330, 81, 21))
        self.label_3.setStyleSheet("background-color:rgb(2, 36, 64); color:rgb(255, 255, 255); font-size:12pt;")
        self.label_3.setObjectName("label_3")
        
        self.kunci_n = QtWidgets.QLineEdit(Enkripsi)
        self.kunci_n.setGeometry(QtCore.QRect(190, 330, 71, 21))
        self.kunci_n.setStyleSheet("color:rgb(255, 255, 255); background-color:rgb(2, 36, 64); font:11pt;")
        self.kunci_n.setObjectName("kunci_n")
        
        self.kunci_e = QtWidgets.QLineEdit(Enkripsi)
        self.kunci_e.setGeometry(QtCore.QRect(280, 330, 71, 20))
        self.kunci_e.setStyleSheet("color:rgb(255, 255, 255); background-color:rgb(2, 36, 64); font-size:11pt;")
        self.kunci_e.setObjectName("kunci_e")
        
        self.encryptbutton = QtWidgets.QPushButton(Enkripsi)
        self.encryptbutton.setGeometry(QtCore.QRect(240, 400, 75, 31))
        self.encryptbutton.setStyleSheet("background-color:rgb(255, 255, 255); border-radius: 5px")
        self.encryptbutton.setObjectName("encryptbutton")
        
        self.savebutton = QtWidgets.QPushButton(Enkripsi)
        self.savebutton.setGeometry(QtCore.QRect(350, 400, 75, 31))
        self.savebutton.setStyleSheet("background-color:rgb(255, 255, 255); border-radius:5px;")
        self.savebutton.setObjectName("savebutton")
        
        self.label_5 = QtWidgets.QLabel(Enkripsi)
        self.label_5.setGeometry(QtCore.QRect(0, 0, 600, 700))
        self.label_5.setStyleSheet("background-image:url(:/newPrefix/Users/user/Downloads/Free Vector _ Gradient geometric blue technology background.jpeg);")
        self.label_5.setText("")
        self.label_5.setPixmap(QtGui.QPixmap("../../../../Users/user/Downloads/Premium Vector _ Circuit board Motherboard_ Blue technology background.jpeg"))
        self.label_5.setScaledContents(True)
        self.label_5.setObjectName("label_5")
        
        self.label_4 = QtWidgets.QLabel(Enkripsi)
        self.label_4.setGeometry(QtCore.QRect(220, 110, 181, 31))
        self.label_4.setStyleSheet("color:rgb(255, 255, 255); font-size:15pt;")
        self.label_4.setObjectName("label_4")
        
        self.backbutton = QtWidgets.QPushButton(Enkripsi)
        self.backbutton.setGeometry(QtCore.QRect(30, 630, 81, 31))
        self.backbutton.setStyleSheet("font-size:11pt")
        self.backbutton.setObjectName("backbutton")
        
        self.label_5.raise_()
        self.browsebutton.raise_()
        self.label_3.raise_()
        self.encryptbutton.raise_()
        self.inputdata.raise_()
        self.savebutton.raise_()
        self.label_2.raise_()
        self.kunci_e.raise_()
        self.kunci_n.raise_()
        self.label_4.raise_()
        self.backbutton.raise_()

        self.retranslateUi(Enkripsi)
        QtCore.QMetaObject.connectSlotsByName(Enkripsi)

        if public_key_n is not None and public_key_e is not None:
            self.kunci_n.setText(str(public_key_n))
            self.kunci_e.setText(str(public_key_e))

        # Hubungkan tombol dengan fungsi
        self.encryptbutton.clicked.connect(self.encrypt_data)
        self.savebutton.clicked.connect(self.save_encrypted_data)
        self.browsebutton.clicked.connect(self.browse_file)
        self.backbutton.clicked.connect(self.open_RSA)

    def retranslateUi(self, Enkripsi):
        _translate = QtCore.QCoreApplication.translate
        Enkripsi.setWindowTitle(_translate("Enkripsi", "Dialog"))
        self.label_2.setText(_translate("Enkripsi", "Pilih Data"))
        self.browsebutton.setText(_translate("Enkripsi", "Browse"))
        self.label_3.setText(_translate("Enkripsi", "Input Kunci"))
        self.encryptbutton.setText(_translate("Enkripsi", "Enkripsi"))
        self.savebutton.setText(_translate("Enkripsi", "Simpan"))
        self.label_4.setText(_translate("Enkripsi", "ENKRIPSI DATA RSA"))
        self.backbutton.setText(_translate("Enkripsi", "<-back"))

    def encrypt_data(self):
     try:
        # Step 1: Ambil input data dan kunci dari UI
        data = self.inputdata.text().strip()  # Hapus spasi di awal dan akhir
        if not data:
            QtWidgets.QMessageBox.warning(None, "Peringatan", "Data untuk dienkripsi kosong!")
            return

         # Step 2: Validasi dan ambil nilai kunci publik RSA (n dan e)
        try:
            n = int(self.kunci_n.text())  # Modulus
            e = int(self.kunci_e.text())  # Eksponen publik
        except ValueError:
            QtWidgets.QMessageBox.warning(None, "Peringatan", "Kunci harus berupa bilangan bulat!")
            return

        # Debugging: Print nilai n dan e
        print(f"n: {n}, e: {e}")

        # Buat kunci publik RSA
        try:
            public_key = RSA.construct((n, e))
        except ValueError as ve:
            QtWidgets.QMessageBox.warning(None, "Peringatan", f"Kesalahan saat membuat kunci publik: {ve}")
            return

        # Step 4: Buat cipher RSA dengan padding OAEP
        cipher = PKCS1_OAEP.new(public_key) 

        # Step 5: Inisialisasi proses enkripsi
        encrypted_data = b'' # Variabel untuk menyimpan data terenkripsi
        max_chunk_size = public_key.size_in_bytes() - 42  # Ukuran maksimum data yang dapat dienkripsi dalam satu blok

        # Debugging: Print ukuran chunk maksimum
        print(f"Max chunk size: {max_chunk_size}")

        start_time = time.time()  # Catat waktu mulai

        # Step 6: Enkripsi data per chunk
        data_bytes = data.encode()  # Ubah data menjadi bytes
        for i in range(0, len(data_bytes), max_chunk_size):  # Loop dengan ukuran chunk dalam bytes
            chunk = data_bytes[i:i + max_chunk_size]  # Ambil chunk data

            # Debugging: Print chunk yang akan dienkripsi
            print(f"Encrypting chunk: {chunk}")

            try:
                encrypted_chunk = cipher.encrypt(chunk)  # Enkripsi chunk
                encrypted_data += encrypted_chunk  # Gabungkan hasil enkripsi
            except ValueError as ve:
                QtWidgets.QMessageBox.warning(None, "Peringatan", f"Kesalahan saat mengenkripsi chunk: {ve}")
                return

        end_time = time.time()  # Catat waktu selesai
        elapsed_time_ms = (end_time - start_time) * 1000  # Hitung waktu proses dalam milidetik



        # Encode hasil enkripsi menjadi base64 untuk memastikan kompatibilitas saat menyimpan
        self.encrypted_data = base64.b64encode(encrypted_data).decode()

        # Debugging: Tampilkan hasil enkripsi di konsol
        print(f"Encrypted Data: {self.encrypted_data}")
        print(f"Running time: {elapsed_time_ms:.2f} ms")
        # Step 7: Tampilkan pesan sukses
        QtWidgets.QMessageBox.information(None, "Sukses", "Data berhasil dienkripsi!")
     except Exception as e:
         # Step 8: Tangani error yang tidak terduga
        QtWidgets.QMessageBox.critical(None, "Error", f"Terjadi kesalahan: {e}")



    def save_encrypted_data(self):
        if hasattr(self, 'encrypted_data'):  # Mengecek apakah objek memiliki atribut 'encrypted_data'
            options = QtWidgets.QFileDialog.Options()
            # Menggunakan dialog penyimpanan file untuk memilih nama dan lokasi
            file_name, _ = QtWidgets.QFileDialog.getSaveFileName(None, "Simpan Encrypted Data", "", "Word Documents (*.docx);;All Files (*)", options=options)
            if file_name:
                doc = docx.Document() # Membuat objek dokumen Word baru menggunakan python-docx
                doc.add_paragraph(self.encrypted_data) # Menambahkan data terenkripsi sebagai paragraf dalam dokumen
                doc.save(file_name)  # Menyimpan dokumen ke file yang dipilih

                QtWidgets.QMessageBox.information(None, "Sukses", "Data berhasil disimpan!")  # informasi bahwa data berhasil disimpan
        else:
            QtWidgets.QMessageBox.warning(None, "Peringatan", "Tidak ada data terenkripsi untuk disimpan.")
           # pesan peringatan jika tidak ada data terenkripsi

    def browse_file(self):
        options = QtWidgets.QFileDialog.Options()
        file_name, _ = QtWidgets.QFileDialog.getOpenFileName(None, "Pilih File DOCX", "C:/Abdiel", "Word Documents (*.docx); All Files (*)", options=options)
        if file_name:
            try:
                doc = docx.Document(file_name)
                text = []
                for para in doc.paragraphs:
                    text.append(para.text.strip())
                
                if not any(text):
                    QtWidgets.QMessageBox.warning(None, "Peringatan", "File DOCX kosong!")
                    return

                self.inputdata.setText('\n'.join(text))
            except Exception as e:
                QtWidgets.QMessageBox.critical(None, "Error", f"Terjadi kesalahan saat membaca file: {e}")

    def open_RSA(self):
        """Fungsi untuk membuka jendela start"""
        from RSA import Ui_RSA
        self.window = QtWidgets.QDialog()
        self.ui = Ui_RSA()
        self.ui.setupUi(self.window)
        self.window.show()

if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    Enkripsi = QtWidgets.QDialog()
    ui = Ui_Enkripsi()
    ui.setupUi(Enkripsi)
    Enkripsi.show()
    sys.exit(app.exec_())
