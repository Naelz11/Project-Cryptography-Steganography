# -*- coding: utf-8 -*-
from PyQt5 import QtCore, QtGui, QtWidgets
from elgamal import my_elgamal as elgamal
import docx
import random
import logging
import time

import bgenkripelgamal_rc

class Ui_enkripsielgamal(object):
    def setupUi(self, Dialog, keys=None):
        Dialog.setObjectName("Dialog")
        Dialog.resize(600, 700)
        self.label_5 = QtWidgets.QLabel(Dialog)
        self.label_5.setGeometry(QtCore.QRect(0, 0, 600, 700))
        self.label_5.setStyleSheet("background-image:url(:/newPrefix/Users/user/Downloads/Free Vector _ Gradient geometric blue technology background.jpeg);")
        self.label_5.setText("")
        self.label_5.setPixmap(QtGui.QPixmap("../../../../Users/user/Downloads/Premium Vector _ Circuit board Motherboard_ Blue technology background.jpeg"))
        self.label_5.setScaledContents(True)
        self.label_5.setObjectName("label_5")
        self.label = QtWidgets.QLabel(Dialog)
        self.label.setGeometry(QtCore.QRect(180, 110, 271, 31))
        self.label.setStyleSheet("color:rgb(255, 255, 255); font-size:15pt;")
        self.label.setObjectName("label")
        self.label_2 = QtWidgets.QLabel(Dialog)
        self.label_2.setGeometry(QtCore.QRect(100, 240, 81, 16))
        self.label_2.setStyleSheet("color:rgb(255, 255, 255); font-size:12pt;")
        self.label_2.setObjectName("label_2")
        self.inputdata = QtWidgets.QLineEdit(Dialog)
        self.inputdata.setGeometry(QtCore.QRect(200, 230, 201, 31))
        self.inputdata.setStyleSheet("background-color:rgb(3, 41, 77); color:rgb(255, 255, 255); font-size:11pt")
        self.inputdata.setObjectName("inputdata")
        self.browsebutton = QtWidgets.QPushButton(Dialog)
        self.browsebutton.setGeometry(QtCore.QRect(424, 230, 81, 31))
        self.browsebutton.setStyleSheet("font-size:10pt;")
        self.browsebutton.setObjectName("browsebutton")
        self.label_3 = QtWidgets.QLabel(Dialog)
        self.label_3.setGeometry(QtCore.QRect(100, 340, 91, 21))
        self.label_3.setStyleSheet("color:rgb(255, 255, 255); font-size:12pt;")
        self.label_3.setObjectName("label_3")
        self.kunci_p = QtWidgets.QLineEdit(Dialog)
        self.kunci_p.setGeometry(QtCore.QRect(200, 340, 71, 21))
        self.kunci_p.setStyleSheet("background-color:rgb(3, 41, 77); color:rgb(255, 255, 255); font-size:11pt")
        self.kunci_p.setObjectName("kunci_p")
        self.kunci_g = QtWidgets.QLineEdit(Dialog)
        self.kunci_g.setGeometry(QtCore.QRect(290, 340, 71, 21))
        self.kunci_g.setStyleSheet("background-color:rgb(3, 41, 77); color:rgb(255, 255, 255); font-size:11pt")
        self.kunci_g.setObjectName("kunci_g")
        self.kunci_y = QtWidgets.QLineEdit(Dialog)
        self.kunci_y.setGeometry(QtCore.QRect(380, 340, 71, 21))
        self.kunci_y.setStyleSheet("background-color:rgb(3, 41, 77); color:rgb(255, 255, 255); font-size:11pt")
        self.kunci_y.setObjectName("kunci_y")
        self.encryptbutton = QtWidgets.QPushButton(Dialog)
        self.encryptbutton.setGeometry(QtCore.QRect(240, 400, 81, 31))
        self.encryptbutton.setStyleSheet("")
        self.encryptbutton.setObjectName("encryptbutton")
        self.savebutton = QtWidgets.QPushButton(Dialog)
        self.savebutton.setGeometry(QtCore.QRect(350, 400, 75, 31))
        self.savebutton.setStyleSheet("")
        self.savebutton.setObjectName("savebutton")
        self.backbutton = QtWidgets.QPushButton(Dialog)
        self.backbutton.setGeometry(QtCore.QRect(30, 630, 81, 31))
        self.backbutton.setStyleSheet("font:11pt\n")
        self.backbutton.setObjectName("backbutton")

        self.retranslateUi(Dialog)
        QtCore.QMetaObject.connectSlotsByName(Dialog)
        self.backbutton.clicked.connect(self.open_startwindow)
        self.encryptbutton.clicked.connect(self.encrypt_message)
        self.browsebutton.clicked.connect(self.browse_button)
        self.savebutton.clicked.connect(self.save_button)

        if keys:
            self.kunci_p.setText(str(keys['p']))
            self.kunci_g.setText(str(keys['g']))
            self.kunci_y.setText(str(keys['h']))

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "Dialog"))
        self.label.setText(_translate("Dialog", "ENKRIPSI DATA EL-GAMAL"))
        self.label_2.setText(_translate("Dialog", "Pilih Data"))
        self.browsebutton.setText(_translate("Dialog", "Browse"))
        self.label_3.setText(_translate("Dialog", "Input Kunci"))
        self.encryptbutton.setText(_translate("Dialog", "Enkripsi"))
        self.savebutton.setText(_translate("Dialog", "Simpan"))
        self.backbutton.setText(_translate("Dialog", "<-back"))

    def open_startwindow(self):
        """Fungsi untuk membuka jendela start"""
        from startwindow import Ui_startwindow
        self.window = QtWidgets.QDialog()
        self.ui = Ui_startwindow()
        self.ui.setupUi(self.window)
        self.window.show()

    def encrypt_message(self):
      try:
        # Ambil input dari field yang sesuai
        self.message = self.inputdata.text()
        self.p = int(self.kunci_p.text())  # Nilai p (bilangan prima) untuk enkripsi
        self.g = int(self.kunci_g.text())  # Nilai g (generator) untuk enkripsi
        self.y = int(self.kunci_y.text())  # Nilai y (kunci publik) untuk enkripsi

        # Pastikan field tidak kosong dan valid
        if not self.message or not self.p or not self.g or not self.y:
            raise ValueError("Semua field kunci harus diisi dan valid.")
        
        # List untuk menyimpan hasil enkripsi
        self.encrypted_message = []

        # Mulai hitung waktu enkripsi
        start_time = time.time()

        # Enkripsi setiap karakter dalam pesan
        for char in self.message:
            m = ord(char)  # Ubah karakter menjadi nilai ASCII

            # Pilih nilai acak k
            k = random.randint(1, self.p - 2) # Pilih k secara acak dari rentang 1 hingga p-2

            # Enkripsi nilai a dan b menggunakan pow untuk modular exponentiation yang lebih cepat
            a = pow(self.g, k, self.p)  # Hitung a = g^k mod p
            b = (pow(self.y, k, self.p) * m) % self.p  # Hitung b = (y^k * m) mod p
            
            # Simpan pasangan (a, b)
            self.encrypted_message.append((a, b))  

            # Hitung waktu yang diambil untuk proses enkripsi
            end_time = time.time()
            encryption_time = (end_time - start_time) * 1000 # Waktu enkripsi dalam milidetik

            # Debugging: Tampilkan hasil enkripsi di konsol
            
            print(f"Running time: {encryption_time:.4f} ms")

        # Proses enkripsi selesai, tampilkan hasil dari pasangan pertama
        encrypted_message_str = f"A: {self.encrypted_message[0][0]}, B: {self.encrypted_message[0][1]}"
        QtWidgets.QMessageBox.information(None, "Sukses", "Data berhasil dienkripsi!")
        QtWidgets.QMessageBox.information(None, "Hasil Enkripsi", f"Pesan terenkripsi: {encrypted_message_str}")
        
      except ValueError as ve:
             QtWidgets.QMessageBox.warning(None, "Input Error", f"Kesalahan input: {str(ve)}")
    
      except Exception as e:
            QtWidgets.QMessageBox.warning(None, "Error", f"Terjadi kesalahan: {str(e)}")

        
    def browse_button(self):
        options = QtWidgets.QFileDialog.Options()
        file_name, _ = QtWidgets.QFileDialog.getOpenFileName(None, "Pilih File DOCX", "C:/Abdiel", "Word Documents (*.docx); All Files (*)", options=options)
        if file_name:
            try:
                doc = docx.Document(file_name)
                text = []
                for para in doc.paragraphs:
                    text.append(para.text.strip())
                
                if not any(text):
                    QtWidgets.QMessageBox.warning(None, "Peringatan", "File DOCX kosong!")
                    return

                self.inputdata.setText('\n'.join(text))
            except Exception as e:
                QtWidgets.QMessageBox.critical(None, "Error", f"Terjadi kesalahan saat membaca file: {e}")
    
    def save_button(self):
        if hasattr(self, 'encrypted_message'):
            options = QtWidgets.QFileDialog.Options()
            file_name, _ = QtWidgets.QFileDialog.getSaveFileName(None, "Simpan Encrypted Data", "", "Word Documents (*.docx);;All Files (*)", options=options)
            
            if file_name:
                try:
                    doc = docx.Document()
                    
                    # Simpan hasil enkripsi ke dalam file DOCX
                    for a, b in self.encrypted_message:
                        encrypted_message_str = f"A: {a}, B: {b}"
                        doc.add_paragraph(encrypted_message_str)
                    doc.save(file_name)
                    
                    QtWidgets.QMessageBox.information(None, "Sukses", "Data berhasil disimpan!")
                except Exception as e:
                    QtWidgets.QMessageBox.critical(None, "Error", f"Terjadi kesalahan saat menyimpan file: {e}")
        else:
            QtWidgets.QMessageBox.warning(None, "Peringatan", "Tidak ada data terenkripsi untuk disimpan.")

    

if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    Dialog = QtWidgets.QDialog()
    ui = Ui_enkripsielgamal()
    ui.setupUi(Dialog)
    Dialog.show()
    sys.exit(app.exec_())
