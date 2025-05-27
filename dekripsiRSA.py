from PyQt5 import QtCore, QtGui, QtWidgets
from Crypto.PublicKey import RSA
from Crypto.Cipher import PKCS1_OAEP
import base64
import docx
import time

from startwindow import Ui_startwindow

class Ui_Dekripsi(object):
    def setupUi(self, Dekripsi):
        Dekripsi.setObjectName("Dekripsi")
        Dekripsi.resize(600, 700)
        
        self.label_5 = QtWidgets.QLabel(Dekripsi)
        self.label_5.setGeometry(QtCore.QRect(0, 0, 600, 700))
        self.label_5.setStyleSheet("background-image:url(:/newPrefix/Users/user/Downloads/Free Vector _ Gradient geometric blue technology background.jpeg);")
        self.label_5.setText("")
        self.label_5.setPixmap(QtGui.QPixmap("../../../../Users/user/Downloads/Premium Vector _ Circuit board Motherboard_ Blue technology background.jpeg"))
        self.label_5.setScaledContents(True)
        self.label_5.setObjectName("label_5")

        self.label = QtWidgets.QLabel(Dekripsi)
        self.label.setGeometry(QtCore.QRect(180, 120, 245, 41))
        self.label.setStyleSheet("font: 75 20pt \"MS Shell Dlg 2\"; color:rgb(255, 255, 255)")
        self.label.setObjectName("label")
        
        self.label_2 = QtWidgets.QLabel(Dekripsi)
        self.label_2.setGeometry(QtCore.QRect(100, 240, 81, 21))
        self.label_2.setStyleSheet("color:rgb(255, 255, 255); font-size:12pt;")
        self.label_2.setObjectName("label_2")
        
        self.label_3 = QtWidgets.QLabel(Dekripsi)
        self.label_3.setGeometry(QtCore.QRect(100, 300, 131, 21))
        self.label_3.setStyleSheet("color:rgb(255, 255, 255); font-size:12pt;")
        self.label_3.setObjectName("label_3")
        
        self.label_4 = QtWidgets.QLabel(Dekripsi)
        self.label_4.setGeometry(QtCore.QRect(100, 360, 131, 21))
        self.label_4.setStyleSheet("color:rgb(255, 255, 255); font-size:12pt;")
        self.label_4.setObjectName("label_4")
        
        self.inputdata = QtWidgets.QLineEdit(Dekripsi)
        self.inputdata.setGeometry(QtCore.QRect(210, 230, 201, 31))
        self.inputdata.setStyleSheet("background-color:rgb(3, 41, 77); color:rgb(255, 255, 255); font-size:11pt")
        self.inputdata.setObjectName("inputdata")
        
        self.browsebutton = QtWidgets.QPushButton(Dekripsi)
        self.browsebutton.setGeometry(QtCore.QRect(424, 230, 81, 31))
        self.browsebutton.setStyleSheet("font-size:10pt;")
        self.browsebutton.setObjectName("browsebutton")
        
        self.kunci_d = QtWidgets.QLineEdit(Dekripsi)
        self.kunci_d.setGeometry(QtCore.QRect(240, 300, 90, 31))
        self.kunci_d.setStyleSheet("background-color:rgb(3, 41, 77); color:rgb(255, 255, 255); font-size:11pt")
        self.kunci_d.setObjectName("kunci_d")

        self.modulus_n = QtWidgets.QLineEdit(Dekripsi)
        self.modulus_n.setGeometry(QtCore.QRect(240, 360, 90, 31))
        self.modulus_n.setStyleSheet("background-color:rgb(3, 41, 77); color:rgb(255, 255, 255); font-size:11pt")
        self.modulus_n.setObjectName("modulus_n")
        
        self.decryptbutton = QtWidgets.QPushButton(Dekripsi)
        self.decryptbutton.setGeometry(QtCore.QRect(240, 424, 81, 31))
        self.decryptbutton.setStyleSheet("")
        self.decryptbutton.setObjectName("decryptbutton")
        
        self.savebutton = QtWidgets.QPushButton(Dekripsi)
        self.savebutton.setGeometry(QtCore.QRect(350, 424, 81, 31))
        self.savebutton.setStyleSheet("")
        self.savebutton.setObjectName("savebutton")
        
        self.backbutton = QtWidgets.QPushButton(Dekripsi)
        self.backbutton.setGeometry(QtCore.QRect(30, 630, 81, 31))
        self.backbutton.setStyleSheet("font:11pt")
        self.backbutton.setObjectName("backbutton")

        self.retranslateUi(Dekripsi)
        QtCore.QMetaObject.connectSlotsByName(Dekripsi)

        # Hubungkan tombol dengan fungsi
        self.decryptbutton.clicked.connect(self.decrypt_data)
        self.savebutton.clicked.connect(self.save_decrypted_data)
        self.browsebutton.clicked.connect(self.browse_file)
        self.backbutton.clicked.connect(self.open_RSA)

    def retranslateUi(self, Dekripsi):
        _translate = QtCore.QCoreApplication.translate
        Dekripsi.setWindowTitle(_translate("Dekripsi", "Dialog"))
        self.label.setText(_translate("Dekripsi", "DEKRIPSI DATA RSA"))
        self.label_2.setText(_translate("Dekripsi", "Pilih Data"))
        self.label_3.setText(_translate("Dekripsi", "Input Kunci Privat"))
        self.label_4.setText(_translate("Dekripsi", "Input Modulus (n)"))
        self.browsebutton.setText(_translate("Dekripsi", "Browse"))
        self.decryptbutton.setText(_translate("Dekripsi", "Dekripsi"))
        self.savebutton.setText(_translate("Dekripsi", "Simpan"))
        self.backbutton.setText(_translate("Dekripsi", "<-back"))

    def decrypt_data(self):
      try:
        encrypted_data = self.inputdata.text().strip()
        if not encrypted_data:
            QtWidgets.QMessageBox.warning(None, "Peringatan", "Data terenkripsi kosong!")
            return

        d_str = self.kunci_d.text().strip()
        n_str = self.modulus_n.text().strip()

        if not d_str or not n_str:
            QtWidgets.QMessageBox.warning(None, "Peringatan", "Kunci privat atau modulus kosong!")
            return

        d = int(d_str)
        n = int(n_str)

        start_time = time.time()

        # Buat kunci privat RSA
        private_key = RSA.construct((n, 65537, d))  # 65537 adalah eksponen publik default
        cipher = PKCS1_OAEP.new(private_key)

        # Decode base64 dan dekripsi data
        encrypted_data_bytes = base64.b64decode(encrypted_data)
        decrypted_data = b''
        max_chunk_size = private_key.size_in_bytes()

        # Loop melalui potongan-potongan data terenkripsi
        for i in range(0, len(encrypted_data_bytes), max_chunk_size):
            chunk = encrypted_data_bytes[i:i+max_chunk_size]
            decrypted_data += cipher.decrypt(chunk)

        self.decrypted_data = decrypted_data.decode()

        QtWidgets.QMessageBox.information(None, "Sukses", "Data berhasil didekripsi!")
      except Exception as e:
        QtWidgets.QMessageBox.critical(None, "Error", f"Terjadi kesalahan: {e}")

      end_time = time.time()
      elapsed_time_ms = (end_time - start_time) * 1000  # Konversi ke milidetik
      print(f"Running time: {elapsed_time_ms:.2f} ms")
        
    def save_decrypted_data(self):
        if hasattr(self, 'decrypted_data'):
            options = QtWidgets.QFileDialog.Options()
            file_name, _ = QtWidgets.QFileDialog.getSaveFileName(None, "Simpan Decrypted Data", "", "Word Documents (*.docx)", options=options)
            if file_name:
                doc = docx.Document()
                doc.add_paragraph(self.decrypted_data)
                doc.save(file_name)
                QtWidgets.QMessageBox.information(None, "Sukses", "Data berhasil disimpan!")
        else:
            QtWidgets.QMessageBox.warning(None, "Peringatan", "Tidak ada data terdekripsi untuk disimpan.")

    def browse_file(self):
        options = QtWidgets.QFileDialog.Options()
        file_name, _ = QtWidgets.QFileDialog.getOpenFileName(None, "Pilih File DOCX", "C:/Abdiel", "Word Documents (*.docx); All Files (*)", options=options)
        if file_name:
            try:
                doc = docx.Document(file_name)
                text = []
                for para in doc.paragraphs:
                    text.append(para.text.strip())
                
                if not any(text):
                    QtWidgets.QMessageBox.warning(None, "Peringatan", "File DOCX kosong!")
                    return

                self.inputdata.setText('\n'.join(text))
            except Exception as e:
                QtWidgets.QMessageBox.critical(None, "Error", f"Terjadi kesalahan saat membaca file: {e}")

    def open_RSA(self):
        """Fungsi untuk membuka jendela start"""
        from RSA import Ui_RSA
        self.window = QtWidgets.QDialog()
        self.ui = Ui_RSA()
        self.ui.setupUi(self.window)
        self.window.show()


